from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage
from database.db import get_all_data, backup_db
from database.models import UserData
from datetime import datetime
import logging
import os
import tempfile
from pathlib import Path
from PIL import Image

logger = logging.getLogger(__name__)


def resolve_photo_path(photo_path: str) -> str | None:
    """Попробовать разрешить путь до фото: абсолютный, относительный, или в папке photos."""
    if not photo_path:
        return None
    # if already absolute and exists
    if os.path.isabs(photo_path) and os.path.exists(photo_path):
        return photo_path
    # as given (relative to cwd)
    if os.path.exists(photo_path):
        return photo_path
    # try relative to project root
    project_root = Path(__file__).resolve().parent.parent
    candidate = project_root / photo_path
    if candidate.exists():
        return str(candidate)
    # try photos directory
    candidate = project_root / 'photos' / Path(photo_path).name
    if candidate.exists():
        return str(candidate)
    return None

def export_to_excel() -> str:
    """Создает красиво оформленный Excel-файл со встроенными фото"""
    backup_db()
    data = get_all_data()
    
    if not data:
        logger.warning("No data to export")
        return None
    
    # Создаем книгу и лист
    wb = Workbook()
    ws = wb.active
    ws.title = "Данные учреждений"
    
    # Русские заголовки
    headers = [
        'ID', 'Telegram ID', 'ФИО', 'Username', 'Телефон', 'Тип', 
        'Название', 'Адрес', 'Ориентир', 'Широта', 'Долгота', 
        'Фото', 'Дата создания'
    ]
    
    # Добавляем заголовки
    ws.append(headers)
    
    # Стили для заголовков
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="3673A5", end_color="3673A5", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    # Форматируем заголовки
    for col in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col)
        cell.value = headers[col-1]
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = border
    
    # Список временных файлов для удаления после сохранения
    temp_files = []
    
    # Заполняем данными
    for row_idx, record in enumerate(data, start=2):
        ws.append([
            record.id,
            record.telegram_id,
            record.full_name or '',
            record.username or '',
            record.phone_number or '',
            record.institution_type or '',
            record.institution_name or '',
            record.address or '',
            record.landmark or '',
            f"{record.latitude:.4f}" if record.latitude else '',
            f"{record.longitude:.4f}" if record.longitude else '',
            '',  # ← Колонка для фото (оставляем пустой)
            record.created_at.strftime('%d.%m.%Y %H:%M') if record.created_at else ''
        ])
        
        # Форматируем строку данных
        for col in range(1, 14):
            cell = ws.cell(row=row_idx, column=col)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = border
        
        # ✅ ВСТАВЛЯЕМ ФОТО В ЯЧЕЙКУ
        resolved_path = resolve_photo_path(record.photo_path) if record.photo_path else None
        if resolved_path:
            try:
                # Оптимизируем фото для Excel (уменьшаем размер)
                img = Image.open(resolved_path)

                # Уменьшаем до разумного размера (макс 300px по ширине)
                max_width = 300
                if img.width > max_width:
                    ratio = max_width / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)

                # Сохраняем временно с уникальным именем в безопасном tmp-файле
                tf = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
                tf_name = tf.name
                tf.close()
                img.save(tf_name, "JPEG", quality=85)
                temp_files.append(tf_name)  # Запоминаем для удаления

                # Вставляем в Excel
                excel_img = XLImage(tf_name)

                # Масштабируем для ячейки (высота ~80 пикселей)
                excel_img.width = 80
                excel_img.height = 80

                # Привязываем к ячейке L (колонка "Фото")
                cell_address = f"L{row_idx}"
                ws.add_image(excel_img, cell_address)

                logger.info(f"Фото добавлено для записи {record.id} (from {resolved_path})")
            except Exception as e:
                logger.error(f"Ошибка при добавлении фото {resolved_path}: {e}")
                # Если фото не удалось добавить, пишем текст
                ws.cell(row=row_idx, column=12).value = "Ошибка загрузки"
    
    # Автоподгонка ширины колонок
    for column in ws.columns:
        max_length = 0
        column_letter = get_column_letter(column[0].column)
        
        # Для колонки с фото делаем фиксированную ширину
        if column_letter == 'L':
            ws.column_dimensions[column_letter].width = 12
            continue
        
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 20)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Высота строк (увеличиваем для фото)
    ws.row_dimensions[1].height = 20  # Заголовок
    for row_idx in range(2, len(data) + 2):
        # Проверяем, есть ли фото в этой строке
        record = data[row_idx - 2]
        resolved_for_row = resolve_photo_path(record.photo_path) if record.photo_path else None
        if resolved_for_row:
            ws.row_dimensions[row_idx].height = 60  # Высота для фото
        else:
            ws.row_dimensions[row_idx].height = 20  # Обычная высота
    
    # Имя файла с датой
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    file_path = f'data_{timestamp}.xlsx'
    
    # ✅ СОХРАНЯЕМ ФАЙЛ (теперь Excel читает временные файлы)
    wb.save(file_path)
    
    # ✅ УДАЛЯЕМ ВРЕМЕННЫЕ ФАЙЛЫ ПОСЛЕ СОХРАНЕНИЯ
    for temp_file in temp_files:
        try:
            if os.path.exists(temp_file):
                os.remove(temp_file)
                logger.debug(f"Удален временный файл: {temp_file}")
        except Exception as e:
            logger.error(f"Не удалось удалить временный файл {temp_file}: {e}")
    
    logger.info(f"Exported {len(data)} records to {file_path}")
    return file_path

def export_to_word() -> str:
    """Экспорт в Word с фото"""
    backup_db()
    data = get_all_data()
    if not data:
        return None
        
    from docx import Document
    from docx.shared import Inches
    
    doc = Document()
    doc.add_heading('Отчёт СЭС - Учреждения', 0)
    
    for idx, record in enumerate(data, 1):
        doc.add_heading(f'{idx}. {record.institution_name}', level=1)
        
        # Создаём таблицу для каждой записи
        table = doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Заполняем данные
        table.rows[0].cells[0].text = 'Тип'
        table.rows[0].cells[1].text = record.institution_type or ''
        
        table.rows[1].cells[0].text = 'Адрес'
        table.rows[1].cells[1].text = record.address or ''
        
        table.rows[2].cells[0].text = 'Ориентир'
        table.rows[2].cells[1].text = record.landmark or ''
        
        table.rows[3].cells[0].text = 'Координаты'
        table.rows[3].cells[1].text = f"{record.latitude}, {record.longitude}" if record.latitude else ''
        
        table.rows[4].cells[0].text = 'Телефон'
        table.rows[4].cells[1].text = record.phone_number or ''
        
        table.rows[5].cells[0].text = 'ФИО'
        table.rows[5].cells[1].text = record.full_name or ''
        
        table.rows[6].cells[0].text = 'Username'
        table.rows[6].cells[1].text = record.username or ''
        
        table.rows[7].cells[0].text = 'Telegram ID'
        table.rows[7].cells[1].text = str(record.telegram_id)
        
        table.rows[8].cells[0].text = 'Дата создания'
        table.rows[8].cells[1].text = record.created_at.strftime('%d.%m.%Y %H:%M') if record.created_at else ''
        
        # ✅ ДОБАВЛЯЕМ ФОТО
        resolved_path = resolve_photo_path(record.photo_path) if record.photo_path else None
        if resolved_path:
            table.rows[9].cells[0].text = 'Фото'
            try:
                # Вставляем фото в ячейку
                paragraph = table.rows[9].cells[1].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(resolved_path, width=Inches(2.5))
                logger.info(f"Фото добавлено в Word для записи {record.id}")
            except Exception as e:
                logger.error(f"Ошибка при добавлении фото в Word: {e}")
                table.rows[9].cells[1].text = 'Ошибка загрузки фото'
        else:
            table.rows[9].cells[0].text = 'Фото'
            table.rows[9].cells[1].text = 'Не отправлено'
        
        # Добавляем разделитель между записями
        doc.add_paragraph()
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    file_path = f'data_word_{timestamp}.docx'
    doc.save(file_path)
    logger.info(f"Exported {len(data)} records to Word")
    return file_path